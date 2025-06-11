from docx import Document
from app.models import UserDetails, OpportunityParticipant, Opportunity, UserPoints
from app.extensions import db
from datetime import datetime
from io import BytesIO
import requests
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_opportunity_description(opp, user_skills):
    print(f"Generating summary for opportunity: {opp.title}")
    prompt = f"""
    Write a concise, professional CV entry in this format only do not add this (Here is a concise, professional CV entry for the user):

    - Contribution: two lines
    - Skills demonstrated: bullet list

    Details:
    Opportunity Title: {opp.title}
    Description: {opp.description}
    Location: {opp.location}
    Dates: {opp.start_date} to {opp.end_date}
    Skills Highlighted: {', '.join(user_skills)}
    """
    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": "llama3.2:1b",
                "prompt": prompt,
                "stream": False
            }
        )
        response.raise_for_status()
        summary = response.json().get("response", "").strip()
        print("Summary generated successfully.")
        return summary
    except Exception as e:
        print(f"Error generating summary for {opp.title}: {e}")
        return f"Error generating summary: {str(e)}"

def generate_user_cv_word(user_id, file_path="user_cv.docx"):
    print(f"Fetching user with ID: {user_id}")
    user = UserDetails.query.get(user_id)
    if not user:
        print("User not found!")
        raise ValueError("User not found")

    print(f"User found: {user.first_name}")
    user_skills = [skill.name for skill in user.skills]
    print(f"User skills: {user_skills}")

    print("Fetching user participations...")
    participations = (
        db.session.query(Opportunity, UserPoints)
        .join(OpportunityParticipant, Opportunity.id == OpportunityParticipant.opportunity_id)
        .join(UserPoints, (UserPoints.opportunity_id == Opportunity.id) & (UserPoints.user_id == user_id), isouter=True)
        .filter(OpportunityParticipant.user_id == user_id, Opportunity.is_deleted == False)
        .order_by(Opportunity.end_date.desc())  
        .limit(3)  
        .all()
    )
    print(f"Found {len(participations)} participations.")

    doc = Document()
    full_name = f"{user.first_name} {user.last_name}"
    heading = doc.add_heading(f"{full_name}", 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    row = table.rows[0].cells

    row[0].text = f"Phone: {user.phone_number}"
    row[1].text = f"Address: {user.country}, {user.city}"

    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_heading("Participated Opportunities", level=1)

    if not participations:
        doc.add_paragraph("No opportunities participated.")
        print("No opportunities found.")
    else:
        for i, (opp, points_obj) in enumerate(participations[:3], 1):
            print(f"Adding opportunity {i}: {opp.title}")
            doc.add_heading(f"{i}. {opp.title}", level=2)
            doc.add_paragraph(f"{opp.location} ({opp.start_date.strftime('%Y-%m-%d')} to {opp.end_date.strftime('%Y-%m-%d')})")
            summary = generate_opportunity_description(opp, user_skills)
            if summary.lower().startswith("here is") or summary.lower().startswith("sure"):
                summary = "\n".join(summary.splitlines()[1:]).strip()
            doc.add_paragraph(summary)

    if isinstance(file_path, BytesIO):
        print("Saving CV to memory buffer...")
        doc.save(file_path)
    else:
        print(f"Saving CV to file: {file_path}")
        doc.save(file_path)

    print("CV generation complete.")
    return file_path
